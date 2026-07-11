import streamlit as st
import os
import re
from io import BytesIO
from langchain_community.vectorstores import FAISS
from langchain_google_genai import GoogleGenerativeAIEmbeddings, ChatGoogleGenerativeAI
from langchain_classic.chains import create_retrieval_chain
from langchain_classic.chains.combine_documents import create_stuff_documents_chain
from langchain_core.prompts import ChatPromptTemplate
from docx import Document
from fpdf import FPDF

st.set_page_config(page_title="Chatbot Teori Konspirasi", page_icon="🕵️")
st.title("🕵️ Chatbot Edukasi Teori Konspirasi Dunia")
st.caption("Menjelaskan asal-usul teori konspirasi terkenal beserta fakta pembandingnya — untuk edukasi & literasi media.")

INDEX_DIR = "faiss_index"

# ============================
# API KEY
# ============================
GOOGLE_API_KEY = st.sidebar.text_input("Masukkan Google API Key:", type="password")

if not GOOGLE_API_KEY:
    st.warning("Masukkan API Key Gemini dulu di sidebar. Ambil gratis di https://aistudio.google.com/apikey")
    st.stop()

os.environ["GOOGLE_API_KEY"] = GOOGLE_API_KEY

# ============================
# LOAD VECTOR DATABASE
# ============================
if not os.path.exists(INDEX_DIR):
    st.error("Vector database belum ada. Jalankan 'python build_vectorstore.py' dulu di terminal.")
    st.stop()

@st.cache_resource
def load_vectorstore(_api_key):
    embeddings = GoogleGenerativeAIEmbeddings(model="gemini-embedding-001")
    return FAISS.load_local(INDEX_DIR, embeddings, allow_dangerous_deserialization=True)

vectorstore = load_vectorstore(GOOGLE_API_KEY)
retriever = vectorstore.as_retriever(search_kwargs={"k": 4})

# ============================
# TOGGLE MODE: RAG vs UMUM
# ============================
mode = st.sidebar.radio(
    "Mode chatbot:",
    ["RAG (jawab dari dataset konspirasi)", "Umum (jawab bebas, termasuk bikin dokumen)"]
)

llm = ChatGoogleGenerativeAI(model="gemini-2.5-flash", temperature=0.3)

# LLM khusus buat generate RPS: max_output_tokens dinaikkan & temperature diturunin
# biar gak gampang ngaco/ngulang-ngulang (degenerate loop) waktu ngarang teks panjang.
llm_rps = ChatGoogleGenerativeAI(
    model="gemini-2.5-flash",
    temperature=0.2,
    max_output_tokens=4096,
)

rag_system_prompt = """Kamu adalah chatbot edukasi tentang teori konspirasi dunia.
Jelaskan secara objektif berdasarkan konteks berikut:
1. Apa isi teori konspirasi tersebut
2. Dari mana/kapan teori itu muncul
3. Fakta atau bantahan ilmiah terhadap teori tersebut

Selalu bersikap netral dan edukatif, jangan membenarkan teori konspirasi sebagai fakta.

Konteks:
{context}"""

rag_prompt = ChatPromptTemplate.from_messages([
    ("system", rag_system_prompt),
    ("human", "{input}"),
])

document_chain = create_stuff_documents_chain(llm, rag_prompt)
qa_chain = create_retrieval_chain(retriever, document_chain)

# ============================
# GENERATOR RPS (SIDEBAR)
# ============================
st.sidebar.markdown("---")
st.sidebar.subheader("📚 Generator RPS")
jumlah_pertemuan = st.sidebar.number_input("Jumlah pertemuan:", min_value=1, max_value=20, value=16)
generate_rps_button = st.sidebar.button("Buat RPS dari Topik Dataset")

def get_topics_with_content(max_chars=800):
    """Ambil daftar topik BESERTA cuplikan isi aslinya dari folder data/,
    biar RPS yang dibuat beneran grounded ke dataset, bukan cuma nama file."""
    topics = []
    if not os.path.exists("data"):
        return topics
    for fname in sorted(os.listdir("data")):
        if fname.endswith(".txt"):
            topic_name = fname.replace(".txt", "").replace("_", " ").strip()
            path = os.path.join("data", fname)
            try:
                with open(path, "r", encoding="utf-8", errors="ignore") as f:
                    content = f.read().strip()
            except Exception:
                content = ""
            snippet = content[:max_chars]
            topics.append({"nama": topic_name, "ringkasan": snippet})
    return topics

def distribute_topics_to_meetings(topics, jumlah_pertemuan):
    """Bagi topik ke tiap pertemuan secara deterministik di Python (bukan
    diserahin ke LLM), supaya semua topik dataset PASTI kepakai dan jumlah
    pertemuan PASTI sesuai input. Pertemuan pertama = pengantar, satu
    pertemuan di tengah = UTS, pertemuan terakhir = UAS/presentasi.
    """
    schedule = []

    if jumlah_pertemuan <= 2:
        # Kasus ekstrem, minimal ada pengantar & penutup
        schedule.append({"no": 1, "jenis": "pengantar", "topics": []})
        for n in range(2, jumlah_pertemuan + 1):
            schedule.append({"no": n, "jenis": "penutup", "topics": topics})
        return schedule

    uts_index = jumlah_pertemuan // 2  # pertemuan UTS = tengah semester (mis. 16 pertemuan -> UTS di 8)
    uas_index = jumlah_pertemuan       # pertemuan terakhir = UAS/presentasi

    materi_slots = [
        n for n in range(1, jumlah_pertemuan + 1)
        if n not in (1, uts_index, uas_index)
    ]

    schedule.append({"no": 1, "jenis": "pengantar", "topics": []})

    if materi_slots and topics:
        # Bagi topik serata mungkin ke slot materi. Kalau topik lebih banyak
        # dari slot, satu pertemuan bisa dapat >1 topik (dikelompokkan).
        n_slots = len(materi_slots)
        n_topics = len(topics)
        base = n_topics // n_slots
        extra = n_topics % n_slots

        idx = 0
        for i, slot in enumerate(materi_slots):
            take = base + (1 if i < extra else 0)
            take = max(take, 1) if idx < n_topics else 0
            chunk = topics[idx: idx + take]
            idx += take
            schedule.append({"no": slot, "jenis": "materi", "topics": chunk})
        # kalau ada sisa topik (harusnya gak ada, tapi jaga-jaga)
        if idx < n_topics:
            schedule[-1]["topics"].extend(topics[idx:])
    else:
        for slot in materi_slots:
            schedule.append({"no": slot, "jenis": "materi", "topics": []})

    schedule.append({"no": uts_index, "jenis": "uts", "topics": []})
    schedule.append({"no": uas_index, "jenis": "uas", "topics": []})

    schedule.sort(key=lambda x: x["no"])
    return schedule

def looks_repetitive(text, min_len=300):
    """Deteksi kasar kalau output LLM keulang-ulang (tanda loop rusak),
    biar bisa langsung retry daripada nyimpen sampah."""
    if len(text) < min_len:
        return False
    sample = text[:2000]
    for i in range(0, len(sample) - 40, 40):
        chunk = sample[i:i + 40]
        if chunk.strip() and sample.count(chunk) > 6:
            return True
    return False

def call_llm_with_retry(llm_obj, prompt, max_retries=2):
    """Panggil LLM, kalau hasilnya keulang-ulang (loop rusak) atau error,
    coba ulang beberapa kali."""
    last_answer = ""
    for attempt in range(max_retries + 1):
        try:
            response = llm_obj.invoke(prompt)
            answer = response.content or ""
        except Exception as e:
            answer = f"[Error saat generate: {e}]"

        if not looks_repetitive(answer):
            return answer
        last_answer = answer

    return last_answer

def build_rps_header_prompt(jumlah_pertemuan, semua_topik_nama):
    daftar_str = "\n".join(f"- {t}" for t in semua_topik_nama)
    return f"""Kamu ahli pengembangan kurikulum. Tulis DUA bagian ini saja untuk mata kuliah
"Literasi Media dan Teori Konspirasi" ({jumlah_pertemuan} kali pertemuan, SKS 3,
semester genap 2026, program studi Sistem Informasi):

1. Identitas Mata Kuliah (nama, kode, SKS, semester, program studi)
2. Capaian Pembelajaran Mata Kuliah (CPMK), fokus pada literasi media, berpikir
   kritis, dan analisis informasi. Buat 5-7 poin CPMK yang jelas dan spesifik.

Konteks: mata kuliah ini akan membahas topik-topik konspirasi berikut sepanjang
semester (jangan dibahas satu-satu di sini, cukup jadi bahan pertimbangan CPMK):
{daftar_str}

ATURAN FORMAT (WAJIB):
- JANGAN pakai tabel markdown (JANGAN pakai karakter '|').
- JANGAN pakai '---' sebagai pemisah.
- Gunakan teks biasa dengan heading '#' dan bold '**teks**' seperlunya.
- Jangan menulis ulang instruksi ini di jawaban, langsung ke isi."""

def build_rps_batch_prompt(batch_slots):
    """Bikin prompt buat satu batch pertemuan (materi/uts/uas/pengantar),
    lengkap dengan cuplikan isi topik asli biar grounded ke dataset."""
    blocks = []
    for slot in batch_slots:
        if slot["jenis"] == "pengantar":
            blocks.append(
                f"Pertemuan {slot['no']}: Pengantar mata kuliah, kontrak kuliah, "
                f"dan pengenalan konsep dasar literasi media & teori konspirasi."
            )
        elif slot["jenis"] == "uts":
            blocks.append(f"Pertemuan {slot['no']}: Ujian Tengah Semester (UTS) — evaluasi materi pertemuan sebelumnya.")
        elif slot["jenis"] == "uas":
            blocks.append(f"Pertemuan {slot['no']}: Ujian Akhir Semester (UAS) / Presentasi akhir mahasiswa.")
        else:
            topik_detail = "\n".join(
                f"  * {t['nama']}: {t['ringkasan'][:500]}" for t in slot["topics"]
            ) or "  (tidak ada topik khusus)"
            nama_topik = ", ".join(t["nama"] for t in slot["topics"]) or "materi umum"
            blocks.append(
                f"Pertemuan {slot['no']}: Bahan kajian = {nama_topik}\n"
                f"Cuplikan isi materi (gunakan sebagai dasar, JANGAN mengarang di luar ini):\n{topik_detail}"
            )

    daftar_pertemuan_str = "\n\n".join(blocks)

    return f"""Kamu ahli pengembangan kurikulum. Untuk SETIAP pertemuan di bawah ini,
tulis rincian RPS dengan format berikut (per pertemuan, TANPA tabel markdown,
TANPA karakter '|', TANPA '---'):

**Pertemuan [nomor]: [judul singkat]**
- Sub-CPMK: [1 kalimat, spesifik ke topik pertemuan ini]
- Bahan Kajian: [ringkas dari cuplikan yang diberikan, JANGAN generik]
- Metode Pembelajaran: [ceramah/diskusi/studi kasus/dll, pilih yang relevan]
- Indikator Penilaian: [1-2 kalimat, terukur]

Data pertemuan yang harus ditulis (WAJIB semua ditulis, jangan ada yang dilewat,
dan jangan gabungkan lebih dari yang diminta):

{daftar_pertemuan_str}

Langsung tulis isinya per pertemuan, tanpa basa-basi pembuka atau penutup."""

def build_rps_footer_prompt(jumlah_pertemuan):
    return """Tulis DUA bagian ini saja untuk RPS mata kuliah "Literasi Media dan Teori Konspirasi":

1. Kriteria dan Bobot Penilaian (misal: tugas, UTS, UAS, partisipasi, presentasi — dengan persentase yang jumlahnya 100%)
2. Daftar Pustaka (boleh menyertakan Wikipedia sebagai sumber data awal, ditambah
   beberapa referensi akademik tentang media literacy/misinformation/critical thinking)

ATURAN FORMAT (WAJIB):
- JANGAN pakai tabel markdown, JANGAN pakai karakter '|', JANGAN pakai '---'.
- Gunakan teks biasa dengan heading '#' dan bold '**teks**' seperlunya.
- Langsung ke isi, tanpa basa-basi pembuka/penutup."""

def generate_full_rps(jumlah_pertemuan, batch_size=4):
    """Generate RPS lengkap secara bertahap (header -> per-batch pertemuan ->
    footer) supaya grounded ke dataset asli dan gak kena loop/potongan token."""
    topics = get_topics_with_content()
    semua_topik_nama = [t["nama"] for t in topics]
    schedule = distribute_topics_to_meetings(topics, jumlah_pertemuan)

    parts = []

    header_prompt = build_rps_header_prompt(jumlah_pertemuan, semua_topik_nama)
    parts.append(call_llm_with_retry(llm_rps, header_prompt))

    parts.append(f"\n\n# Distribusi Topik dan Rincian Pertemuan ({jumlah_pertemuan} Pertemuan)\n")

    for i in range(0, len(schedule), batch_size):
        batch = schedule[i:i + batch_size]
        batch_prompt = build_rps_batch_prompt(batch)
        batch_result = call_llm_with_retry(llm_rps, batch_prompt)
        parts.append(batch_result)

    footer_prompt = build_rps_footer_prompt(jumlah_pertemuan)
    parts.append(call_llm_with_retry(llm_rps, footer_prompt))

    return "\n\n".join(p.strip() for p in parts if p and p.strip())

# ============================
# HELPER: WRAP KATA PANJANG BIAR PDF GAK ERROR
# ============================
def wrap_long_words(text, chunk_size=20):
    """Potong paksa kata yang lebih dari chunk_size karakter, biar pasti muat di PDF."""
    def breaker(match):
        word = match.group(0)
        return ' '.join(word[i:i+chunk_size] for i in range(0, len(word), chunk_size))
    return re.sub(r'\S{' + str(chunk_size + 1) + r',}', breaker, text)

# ============================
# HELPER: BERSIHIN MARKDOWN BUAT PDF
# ============================
def clean_for_pdf(text):
    """Bersihin format markdown yang bikin PDF berantakan."""
    text = re.sub(r'^\|?[\s\-:|=_]+\|?$', '', text, flags=re.MULTILINE)
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'^#{1,3}\s+', '', text, flags=re.MULTILINE)
    text = text.replace("|", " - ")
    return text

# ============================
# HELPER: PARSING MARKDOWN DASAR BUAT WORD
# ============================
def add_markdown_paragraph(doc, line):
    """Tambah paragraf ke Word dengan parsing markdown dasar (bold, heading, list, garis pemisah)."""
    line = line.rstrip("\n")
    stripped = line.strip()

    if re.match(r'^[\-=_]{3,}$', stripped):
        return

    if stripped == "":
        doc.add_paragraph("")
        return

    heading_match = re.match(r'^(#{1,3})\s+(.*)', stripped)
    if heading_match:
        level = len(heading_match.group(1))
        text = heading_match.group(2)
        doc.add_heading(text, level=level)
        return

    bullet_match = re.match(r'^[\-\*]\s+(.*)', stripped)
    if bullet_match:
        text = bullet_match.group(1)
        paragraph = doc.add_paragraph(style="List Bullet")
        add_bold_runs(paragraph, text)
        return

    numbered_match = re.match(r'^\d+\.\s+(.*)', stripped)
    if numbered_match:
        text = numbered_match.group(1)
        paragraph = doc.add_paragraph(style="List Number")
        add_bold_runs(paragraph, text)
        return

    paragraph = doc.add_paragraph()
    add_bold_runs(paragraph, stripped)

def add_bold_runs(paragraph, text):
    """Parse **teks** jadi run bold di dalam sebuah paragraph Word."""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)

# ============================
# HELPER: RENDER TOMBOL DOWNLOAD
# ============================
def render_download_buttons(answer, msg_index):
    col1, col2 = st.columns(2)

    # Word
    doc = Document()
    for line in answer.split("\n"):
        add_markdown_paragraph(doc, line)
    docx_buffer = BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)

    with col1:
        st.download_button(
            label="📄 Download Word (.docx)",
            data=docx_buffer,
            file_name="hasil_chatbot.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key=f"docx_{msg_index}"
        )

    # PDF
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=11)
    pdf.set_auto_page_break(auto=True, margin=15)

    cleaned_answer = clean_for_pdf(answer)

    for line in cleaned_answer.split("\n"):
        safe_line = line.encode("latin-1", "replace").decode("latin-1")
        safe_line = wrap_long_words(safe_line)

        pdf.set_x(pdf.l_margin)  # <-- PENTING: reset posisi X tiap baris

        if safe_line.strip() == "":
            pdf.ln(4)
        else:
            pdf.multi_cell(0, 8, safe_line)

    pdf_bytes = bytes(pdf.output())
    pdf_buffer = BytesIO(pdf_bytes)

    with col2:
        st.download_button(
            label="📕 Download PDF (.pdf)",
            data=pdf_buffer,
            file_name="hasil_chatbot.pdf",
            mime="application/pdf",
            key=f"pdf_{msg_index}"
        )

# ============================
# SESSION STATE
# ============================
if "messages" not in st.session_state:
    st.session_state.messages = []

# ============================
# HANDLE TOMBOL GENERATE RPS
# ============================
if generate_rps_button:
    with st.spinner("Menyusun RPS berdasarkan isi dataset (bisa beberapa saat, digenerate bertahap)..."):
        try:
            answer = generate_full_rps(jumlah_pertemuan)
            if not answer.strip():
                answer = "Maaf, gagal menyusun RPS (hasil kosong). Coba klik lagi."
        except Exception as e:
            answer = f"Maaf, terjadi error: {e}"

    n_topik = len(get_topics_with_content())
    st.session_state.messages.append({"role": "user", "content": f"[Generate RPS dari {n_topik} topik dataset, {jumlah_pertemuan} pertemuan]"})
    st.session_state.messages.append({"role": "assistant", "content": answer})

# ============================
# RENDER RIWAYAT CHAT (SEMUA PESAN, TERMASUK DARI TOMBOL RPS)
# ============================
for i, msg in enumerate(st.session_state.messages):
    with st.chat_message(msg["role"]):
        st.write(msg["content"])
        if msg["role"] == "assistant":
            render_download_buttons(msg["content"], i)

# ============================
# CHAT INPUT
# ============================
user_input = st.chat_input("Tanya soal teori konspirasi, atau minta buatkan dokumen...")

if user_input:
    st.session_state.messages.append({"role": "user", "content": user_input})
    with st.chat_message("user"):
        st.write(user_input)

    with st.spinner("Memproses..."):
        try:
            if mode == "RAG (jawab dari dataset konspirasi)":
                result = qa_chain.invoke({"input": user_input})
                answer = result["answer"]
            else:
                general_prompt = f"""Kamu adalah asisten AI yang membantu berbagai kebutuhan pengguna, termasuk membuat dokumen (RPS, laporan, ringkasan, dll).
Jawab langsung dengan hasil lengkap berdasarkan asumsi yang wajar, jangan banyak bertanya balik kecuali benar-benar penting.
Jawabanmu akan otomatis dikonversi jadi file Word dan PDF oleh sistem, jadi cukup fokus menghasilkan teks yang lengkap dan terstruktur.

Permintaan pengguna: {user_input}"""
                response = llm.invoke(general_prompt)
                answer = response.content
        except Exception as e:
            answer = f"Maaf, terjadi error: {e}"

    with st.chat_message("assistant"):
        st.write(answer)
        render_download_buttons(answer, len(st.session_state.messages))

    st.session_state.messages.append({"role": "assistant", "content": answer})